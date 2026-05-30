"""
Build IEEE two-column DOCX from SybilShield-Core_IEEE_Paper_V4.md.

Layout (matching V2.docx):
  - Page: 8.5x11", margins L/R=0.625", T=0.75", B=1.0"
  - Section 1 (continuous): title + authors + abstract, FULL WIDTH (1 col)
  - Section 2: body text + references, TWO COLUMNS, 720-twip gutter
  - Fonts / sizes exactly as V2: Times New Roman throughout
      title 24pt bold center | author lines 9pt center
      abstract 9pt justified | heading1 10pt bold small-caps center
      heading2 10pt italic   | body 10pt justified, 0.2" first-line
      equations centered italic | references 9pt hanging-indent
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ────────────────────────────────────────────────────────────────

def add_inline(para, text, size=10, font="Times New Roman"):
    """Render **bold**, *italic*, `mono` markup into runs."""
    for part in re.split(r'(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`)', text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2]); run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1]); run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"; run.font.size = Pt(size - 1); continue
        else:
            run = para.add_run(part)
        run.font.name = font
        run.font.size = Pt(size)


def _pg_sectPr(top=1080, right=900, bottom=1440, left=900,
               header=720, footer=720, cols=1, space=720):
    """Return a fully specified <w:sectPr> element."""
    sp = OxmlElement("w:sectPr")
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), "12240"); pgSz.set(qn("w:h"), "15840")
    sp.append(pgSz)
    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:top"),    str(top))
    pgMar.set(qn("w:right"),  str(right))
    pgMar.set(qn("w:bottom"), str(bottom))
    pgMar.set(qn("w:left"),   str(left))
    pgMar.set(qn("w:header"), str(header))
    pgMar.set(qn("w:footer"), str(footer))
    pgMar.set(qn("w:gutter"), "0")
    sp.append(pgMar)
    colsEl = OxmlElement("w:cols")
    if cols > 1:
        colsEl.set(qn("w:num"), str(cols))
    colsEl.set(qn("w:space"), str(space))
    sp.append(colsEl)
    return sp


def apply_sectPr(doc_section, cols=2):
    """Overwrite the existing sectPr of a python-docx Section object."""
    new_sp = _pg_sectPr(cols=cols)
    old_sp = doc_section._sectPr
    old_sp.getparent().replace(old_sp, new_sp)
    doc_section._sectPr = new_sp


def insert_continuous_break(doc, cols_for_this_section=1):
    """
    Append an empty paragraph whose pPr carries the sectPr that ENDS
    the current section.  Word interprets this as a continuous section break.
    cols_for_this_section = column count of the section ending here.
    """
    p_elem = doc.add_paragraph()
    p_elem.paragraph_format.space_before = Pt(0)
    p_elem.paragraph_format.space_after  = Pt(0)
    pPr = p_elem._p.get_or_add_pPr()
    sp = _pg_sectPr(cols=cols_for_this_section)
    typeEl = OxmlElement("w:type")
    typeEl.set(qn("w:val"), "continuous")
    sp.append(typeEl)
    pPr.append(sp)


def add_border_table(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top","left","bottom","right","insideH","insideV"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"),   "single")
                b.set(qn("w:sz"),    "4")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), "000000")
                borders.append(b)
            tcPr.append(borders)


def parse_md_table(lines):
    rows = []
    for ln in lines:
        ln = ln.strip()
        if re.match(r"^\|[-| :]+\|$", ln):
            continue
        rows.append([c.strip() for c in ln.strip("|").split("|")])
    return rows


# ── style setup ─────────────────────────────────────────────────────────────

def setup_styles(doc):
    """Create / update all IEEE-required styles."""
    def get_or_add(name):
        try:
            return doc.styles[name]
        except KeyError:
            return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    # paper title
    s = get_or_add("paper title")
    s.font.name = "Times New Roman"; s.font.size = Pt(24); s.font.bold = True
    s.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after  = Pt(6)
    s.paragraph_format.space_before = Pt(0)

    # Author (affiliation lines)
    s = get_or_add("Author")
    s.font.name = "Times New Roman"; s.font.size = Pt(9)
    s.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after  = Pt(2)

    # Abstract
    s = get_or_add("Abstract")
    s.font.name = "Times New Roman"; s.font.size = Pt(9)
    s.paragraph_format.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.space_before       = Pt(6)
    s.paragraph_format.space_after        = Pt(0)
    s.paragraph_format.first_line_indent  = Inches(0.2)

    # Normal body
    s = doc.styles["Normal"]
    s.font.name = "Times New Roman"; s.font.size = Pt(10)
    s.paragraph_format.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.space_before       = Pt(0)
    s.paragraph_format.space_after        = Pt(0)
    s.paragraph_format.first_line_indent  = Inches(0.2)
    s.paragraph_format.line_spacing       = Pt(12)

    # Heading 1  (section: "I. INTRODUCTION")
    s = doc.styles["Heading 1"]
    s.font.name = "Times New Roman"; s.font.size = Pt(10); s.font.bold = True
    s.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(8)
    s.paragraph_format.space_after  = Pt(4)

    # Heading 2  (subsection: "A. Adversary Definition")
    s = doc.styles["Heading 2"]
    s.font.name = "Times New Roman"; s.font.size = Pt(10)
    s.font.bold = False; s.font.italic = True
    s.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.space_before = Pt(6)
    s.paragraph_format.space_after  = Pt(2)

    # List Paragraph (references)
    s = doc.styles["List Paragraph"]
    s.font.name = "Times New Roman"; s.font.size = Pt(9)
    s.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.left_indent        = Inches(0.3)
    s.paragraph_format.first_line_indent  = Inches(-0.3)
    s.paragraph_format.space_before       = Pt(2)
    s.paragraph_format.space_after        = Pt(2)


# ── main converter ──────────────────────────────────────────────────────────

def convert():
    src = r"C:\IEEE\SybilShield-Core\SybilShield-Core_IEEE_Paper_V4.md"
    dst = r"C:\IEEE\SybilShield-Core\SybilShield-Core_IEEE_Paper_V4.docx"

    with open(src, encoding="utf-8") as fh:
        lines = [l.rstrip("\n") for l in fh]

    doc = Document()
    setup_styles(doc)

    # The document's built-in sectPr will become Section 2 (two-column body).
    # We REPLACE it entirely with our IEEE two-column spec.
    apply_sectPr(doc.sections[0], cols=2)

    in_header          = True   # flips after we insert the section break
    section_break_done = False

    i = 0
    N = len(lines)

    while i < N:
        line = lines[i]

        # blank
        if not line.strip():
            i += 1; continue

        # --- horizontal rule (skip) ---
        if line.strip() == "---":
            i += 1; continue

        # ── TITLE ──────────────────────────────────────────────────────────
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_paragraph(style="paper title")
            add_inline(p, line[2:].strip(), size=24)
            for r in p.runs: r.bold = True
            i += 1; continue

        # ── ABSTRACT / INDEX TERMS ─────────────────────────────────────────
        if line.startswith("*Abstract*") or line.startswith("*Index Terms*"):
            p = doc.add_paragraph(style="Abstract")
            # render "Abstract" keyword in bold-italic
            m = re.match(r"^(\*\w[\w ]*\*)(\s*--\s*)(.*)", line, re.DOTALL)
            if m:
                kw_run = p.add_run(m.group(1)[1:-1])   # strip *
                kw_run.bold = True; kw_run.italic = True
                kw_run.font.name = "Times New Roman"; kw_run.font.size = Pt(9)
                sep_run = p.add_run(m.group(2))
                sep_run.font.name = "Times New Roman"; sep_run.font.size = Pt(9)
                add_inline(p, m.group(3), size=9)
            else:
                add_inline(p, line, size=9)
            i += 1; continue

        # ── AUTHOR NAME LINE  (**1st / 2nd / 3rd Name**) ──────────────────
        if re.match(r"^\*\*\d+(st|nd|rd|th)\s", line):
            display = re.sub(r"^\*\*\d+(st|nd|rd|th)\s+", "", line).rstrip("*")
            p = doc.add_paragraph(style="Author")
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(display)
            r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(10)
            i += 1; continue

        # ── AUTHOR AFFILIATION / EMAIL (plain lines in header region) ─────
        if in_header and not line.startswith("#") and not line.startswith("|") \
                and not line.startswith(">") and not line.startswith("-") \
                and not line.startswith("*") and not re.match(r"^\[\d+\]", line) \
                and not re.match(r"^\d+\.\s", line) and not line.startswith("https://"):
            kw = ["@","Professor","Engineer","Division","College","University",
                  "Inc","USA","India","IEEE","ISACA","Member","Software","CSE",
                  "Department","Malla","HCLTech","Rocket","Dallas","Hyderabad"]
            if any(k in line for k in kw):
                p = doc.add_paragraph(style="Author")
                r = p.add_run(line.strip())
                r.font.name = "Times New Roman"; r.font.size = Pt(9)
                i += 1; continue

        # ── SECTION HEADING  (## …) ────────────────────────────────────────
        if line.startswith("## ") and not line.startswith("### "):
            heading_text = line[3:].strip()

            # First section heading triggers the section break:
            # ends Section 1 (full-width, 1 col) and starts Section 2 (2 col).
            if not section_break_done:
                insert_continuous_break(doc, cols_for_this_section=1)
                section_break_done = True
                in_header = False

            p = doc.add_paragraph(style="Heading 1")
            r = p.add_run(heading_text.upper())
            r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(10)
            i += 1; continue

        # ── SUBSECTION HEADING  (### …) ────────────────────────────────────
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            r = p.add_run(line[4:].strip())
            r.italic = True; r.font.name = "Times New Roman"; r.font.size = Pt(10)
            i += 1; continue

        # ── EQUATION / BLOCK QUOTE  (> …) ─────────────────────────────────
        if line.startswith("> "):
            eq = line[2:].strip().replace("**", "")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(eq); r.italic = True
            r.font.name = "Times New Roman"; r.font.size = Pt(10)
            i += 1; continue

        # ── MARKDOWN TABLE  (| … |) ────────────────────────────────────────
        if line.startswith("|"):
            tbl_lines = []
            while i < N and lines[i].startswith("|"):
                tbl_lines.append(lines[i]); i += 1
            rows = parse_md_table(tbl_lines)
            if rows:
                nc = len(rows[0])
                tbl = doc.add_table(rows=len(rows), cols=nc)
                tbl.style = "Table Grid"
                add_border_table(tbl)
                for ri, row_data in enumerate(rows):
                    for ci, ct in enumerate(row_data):
                        cell = tbl.rows[ri].cells[ci]
                        cell.paragraphs[0].clear()
                        cp = cell.paragraphs[0]
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        add_inline(cp, ct, size=8)
                        if ri == 0:
                            for run in cp.runs: run.bold = True
                gap = doc.add_paragraph()
                gap.paragraph_format.space_after = Pt(4)
            continue

        # ── BULLET LIST  (- …) ────────────────────────────────────────────
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            add_inline(p, line[2:].strip(), size=10)
            i += 1; continue

        # ── NUMBERED LIST  (1. …) ─────────────────────────────────────────
        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            add_inline(p, re.sub(r"^\d+\.\s+", "", line), size=10)
            i += 1; continue

        # ── REFERENCE ENTRY  ([n] …) ──────────────────────────────────────
        if re.match(r"^\[\d+\]", line):
            p = doc.add_paragraph(style="List Paragraph")
            add_inline(p, line, size=9)
            i += 1; continue

        # ── URL ────────────────────────────────────────────────────────────
        if line.startswith("https://"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line.strip())
            r.font.name = "Courier New"; r.font.size = Pt(9)
            i += 1; continue

        # ── BOLD-LABEL PARAGRAPH  (**Label.** body text) ──────────────────
        m = re.match(r"^(\*\*[^*]+\*\*)\s*(.*)", line)
        if m:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.2)
            p.paragraph_format.space_before      = Pt(3)
            p.paragraph_format.space_after       = Pt(3)
            r = p.add_run(m.group(1)[2:-2] + " ")
            r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(10)
            add_inline(p, m.group(2), size=10)
            i += 1; continue

        # ── REGULAR PARAGRAPH ──────────────────────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.2)
        p.paragraph_format.space_before      = Pt(0)
        p.paragraph_format.space_after       = Pt(0)
        p.paragraph_format.line_spacing      = Pt(12)
        add_inline(p, line, size=10)
        i += 1

    doc.save(dst)
    print(f"Saved: {dst}")
    return dst


if __name__ == "__main__":
    convert()
